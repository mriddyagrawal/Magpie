"""Stage 1: summarize a single local file via Moonshot Kimi + PydanticAI."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import sys
from pathlib import Path
from typing import Literal

from dotenv import load_dotenv
from pydantic import BaseModel, Field
from pydantic_ai import Agent, BinaryContent, NativeOutput
from pydantic_ai.models.openai import OpenAIChatModel
from pydantic_ai.providers.openai import OpenAIProvider


IMAGE_EXTS = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
              ".webp": "image/webp", ".gif": "image/gif"}
CODE_EXTS = {".py", ".js", ".ts", ".tsx", ".jsx", ".go", ".rs", ".java",
             ".c", ".cpp", ".h", ".hpp", ".cs", ".rb", ".swift", ".kt",
             ".sh", ".sql", ".json", ".yaml", ".yml", ".toml"}
TEXT_EXTS = {".txt"}
MD_EXTS = {".md", ".markdown"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xlsm"}

MAX_TEXT_CHARS = 120_000

REPO_ROOT = Path(__file__).resolve().parent
SUMMARIES_DIR = REPO_ROOT / "Summaries"

ContentType = Literal["image", "pdf", "docx", "xlsx", "text", "code", "markdown", "other"]

SUPPORTED_EXTS = (
    set(IMAGE_EXTS)
    | PDF_EXTS
    | DOCX_EXTS
    | XLSX_EXTS
    | MD_EXTS
    | TEXT_EXTS
    | CODE_EXTS
)


class SummarizeError(RuntimeError):
    """Raised for per-file failures that should not abort a batch run."""


class FileSummary(BaseModel):
    title: str = Field(description="Short human-readable title for the file (<=80 chars).")
    summary: str = Field(description="Dense 2-5 sentence summary of what the file contains.")
    content_type: ContentType = Field(description="What kind of content the file is.")
    keywords: list[str] = Field(description="3-10 topical keywords/tags for retrieval.")
    key_entities: list[str] = Field(
        description="Named entities: people, orgs, places, products, paths, function names, IDs."
    )


SYSTEM_PROMPT = (
    "You are a file-summarization assistant. Given a single file's content, "
    "produce a FileSummary: a short `title`, a dense 2-5 sentence `summary`, "
    "the `content_type`, 3-10 `keywords`, and `key_entities`. "
    "Be specific and factual. Do not invent content that is not present."
)


def build_agent() -> Agent[None, FileSummary]:
    api_key = os.environ.get("MOONSHOT_API_KEY")
    if not api_key:
        sys.exit("error: MOONSHOT_API_KEY not set (put it in .env)")
    model_name = os.environ.get("MOONSHOT_MODEL", "kimi-k2.5")
    base_url = os.environ.get("MOONSHOT_BASE_URL", "https://api.moonshot.ai/v1")
    model = OpenAIChatModel(
        model_name,
        provider=OpenAIProvider(base_url=base_url, api_key=api_key),
    )
    return Agent(model, output_type=NativeOutput(FileSummary), system_prompt=SYSTEM_PROMPT)


def extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n\n".join(chunks)


def extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            if not any(cell is not None and str(cell).strip() for cell in row):
                continue
            parts.append(",".join("" if cell is None else str(cell) for cell in row))
        parts.append("")
    wb.close()
    return "\n".join(parts)


def build_user_message(path: Path) -> list:
    ext = path.suffix.lower()
    name_hint = f"Filename: {path.name}"

    if ext in IMAGE_EXTS:
        return [
            f"{name_hint}\nSummarize this image.",
            BinaryContent(data=path.read_bytes(), media_type=IMAGE_EXTS[ext]),
        ]

    if ext in PDF_EXTS:
        text = extract_pdf_text(path).strip()
        if not text:
            raise SummarizeError(f"could not extract text from PDF (may be image-only): {path}")
        text = text[:MAX_TEXT_CHARS]
        return [f"{name_hint}\nContent type: pdf\n\n---\n{text}"]

    if ext in DOCX_EXTS:
        text = extract_docx_text(path).strip()
        if not text:
            raise SummarizeError(f"docx appears empty: {path}")
        text = text[:MAX_TEXT_CHARS]
        return [f"{name_hint}\nContent type: docx\n\n---\n{text}"]

    if ext in XLSX_EXTS:
        text = extract_xlsx_text(path).strip()
        if not text:
            raise SummarizeError(f"xlsx appears empty: {path}")
        text = text[:MAX_TEXT_CHARS]
        return [f"{name_hint}\nContent type: xlsx\n\n---\n{text}"]

    if ext in MD_EXTS or ext in TEXT_EXTS or ext in CODE_EXTS:
        ctype = "markdown" if ext in MD_EXTS else ("code" if ext in CODE_EXTS else "text")
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError as e:
            raise SummarizeError(f"file is not valid UTF-8 text: {path}") from e
        text = text[:MAX_TEXT_CHARS]
        return [f"{name_hint}\nContent type: {ctype}\n\n---\n{text}"]

    raise SummarizeError(f"unsupported file type '{ext}' for {path}")


def render_markdown(summary: FileSummary, source_rel: str) -> str:
    keywords = ", ".join(summary.keywords) if summary.keywords else "—"
    entities = ", ".join(summary.key_entities) if summary.key_entities else "—"
    return (
        f"Source: {source_rel}\n\n"
        f"# {summary.title}\n\n"
        f"{summary.summary}\n\n"
        f"**Content type:** {summary.content_type}\n\n"
        f"**Keywords:** {keywords}\n\n"
        f"**Key entities:** {entities}\n"
    )


def hash_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()[:16]


def summary_path_for(path: Path) -> Path:
    return SUMMARIES_DIR / f"{hash_file(path)}.md"


def source_rel_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(REPO_ROOT))
    except ValueError:
        return str(path.resolve())


def write_summary(path: Path, summary: FileSummary) -> Path:
    SUMMARIES_DIR.mkdir(exist_ok=True)
    out_path = summary_path_for(path)
    out_path.write_text(render_markdown(summary, source_rel_path(path)), encoding="utf-8")
    return out_path


def summarize_one(agent: Agent[None, FileSummary], path: Path) -> FileSummary:
    message = build_user_message(path)
    result = agent.run_sync(message)
    write_summary(path, result.output)
    return result.output


def find_supported_files(root: Path) -> list[Path]:
    return sorted(
        p for p in root.rglob("*")
        if p.is_file() and not p.name.startswith(".") and p.suffix.lower() in SUPPORTED_EXTS
    )


def run_batch(agent: Agent[None, FileSummary], root: Path, force: bool) -> None:
    from tqdm import tqdm

    files = find_supported_files(root)
    if not files:
        sys.exit(f"no supported files found under {root}")

    skipped = 0
    errors: list[tuple[Path, str]] = []
    bar = tqdm(files, desc="summarizing", unit="file")
    for path in bar:
        bar.set_postfix_str(path.name[:40])
        try:
            if not force and summary_path_for(path).exists():
                skipped += 1
                continue
            summarize_one(agent, path)
        except SummarizeError as e:
            errors.append((path, str(e)))
            tqdm.write(f"  skip: {e}")
        except Exception as e:  # API errors, decode errors, etc.
            errors.append((path, f"{type(e).__name__}: {e}"))
            tqdm.write(f"  error on {path.name}: {type(e).__name__}: {e}")

    done = len(files) - skipped - len(errors)
    print(f"\ndone: {done} summarized, {skipped} already-cached, {len(errors)} errors")
    if errors:
        print("errors:")
        for p, msg in errors:
            print(f"  - {source_rel_path(p)}: {msg}")


def main() -> None:
    load_dotenv()
    parser = argparse.ArgumentParser(description="Summarize a local file (or directory) via Kimi.")
    parser.add_argument("path", help="File or directory to summarize.")
    parser.add_argument("--force", action="store_true",
                        help="Re-summarize even if a summary already exists for the file's hash.")
    args = parser.parse_args()

    path = Path(args.path)
    if not path.exists():
        sys.exit(f"error: no such path: {path}")

    agent = build_agent()

    if path.is_dir():
        run_batch(agent, path, force=args.force)
        return

    try:
        if not args.force and summary_path_for(path).exists():
            print(f"already summarized: {summary_path_for(path).relative_to(REPO_ROOT)}")
            return
        summary = summarize_one(agent, path)
    except SummarizeError as e:
        sys.exit(f"error: {e}")
    print(json.dumps(summary.model_dump(), indent=2, ensure_ascii=False))
    print(f"\nwrote: {summary_path_for(path).relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
