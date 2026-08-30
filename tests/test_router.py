"""Unit tests for the ingest router.

Covers: tier selection per file type, peek correctness, sensitivity scoring,
T4 cost gates, `.magpieconfig.yaml` overrides, and skip paths.

Fixtures are generated in-test into `tmp_path` — no dependency on the real
repo contents — so these stay hermetic and fast.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from src.router import (
    DEFAULT_T4_BUDGET_MB,
    IMAGE_THUMBNAIL_MIN_DIM,
    IMAGE_THUMBNAIL_SIZE_BYTES,
    T4_MAX_STORAGE_MB_PER_FILE,
    compute_sensitivity_score,
    compute_visual_score,
    decide,
    estimate_t4_cost,
    load_magpieconfig,
    peek,
)


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------

def _write(tmp_path: Path, name: str, content: str | bytes) -> Path:
    p = tmp_path / name
    if isinstance(content, bytes):
        p.write_bytes(content)
    else:
        p.write_text(content, encoding="utf-8")
    return p


def _make_pdf(tmp_path: Path, name: str, pages: list[str]) -> Path:
    """Build a real PDF with the given page texts using pymupdf."""
    import pymupdf
    doc = pymupdf.open()
    for text in pages:
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), text, fontsize=11)
    out = tmp_path / name
    doc.save(str(out))
    doc.close()
    return out


def _make_image_pdf(tmp_path: Path, name: str, n_pages: int = 1) -> Path:
    """Build a PDF with empty pages (no embedded text) — simulates a scan."""
    import pymupdf
    doc = pymupdf.open()
    for _ in range(n_pages):
        doc.new_page()
    out = tmp_path / name
    doc.save(str(out))
    doc.close()
    return out


def _make_docx(tmp_path: Path, name: str, *, paragraphs: int, images: int) -> Path:
    """Build a DOCX with N paragraphs; attach N tiny PNGs to drive image_ratio."""
    from docx import Document
    from PIL import Image

    doc = Document()
    for i in range(paragraphs):
        doc.add_paragraph(f"Paragraph {i}: the quick brown fox jumps over the lazy dog.")
    if images > 0:
        png_path = tmp_path / f"_{name}_img.png"
        Image.new("RGB", (32, 32), "red").save(png_path)
        for _ in range(images):
            doc.add_picture(str(png_path))
    out = tmp_path / name
    doc.save(str(out))
    return out


def _make_image(tmp_path: Path, name: str, size: tuple[int, int]) -> Path:
    from PIL import Image
    p = tmp_path / name
    Image.new("RGB", size, "white").save(p)
    return p


def _sensitive_text() -> str:
    """Text that trips multiple sensitivity detectors. Mimics a bank statement snippet."""
    return (
        "Chase Bank Statement\n"
        "Account: ****1234\n"
        "Period: 2026-03-01 to 2026-03-31\n\n"
        "Transactions:\n"
        "2026-03-02 WHOLE FOODS $127.43\n"
        "2026-03-05 SHELL GAS $52.19\n"
        "2026-03-10 NETFLIX SUBSCRIPTION $15.99\n"
        "2026-03-15 PAYCHECK DIRECT DEPOSIT $3,420.00\n"
        "2026-03-22 AMAZON.COM $89.50\n\n"
        "Subtotal: $285.11\n"
        "Balance: $2,134.89\n"
        "Amount due: $0.00\n"
    )


# ---------------------------------------------------------------------------
# Peek — basic shape per file type
# ---------------------------------------------------------------------------

def test_peek_text_small(tmp_path: Path):
    p = _write(tmp_path, "note.md", "# hello\n\nSome notes here.\n")
    r = peek(p)
    assert r.ext == ".md"
    assert r.extractable is True
    assert r.size_bytes > 0
    assert r.peek_error is None


def test_peek_code(tmp_path: Path):
    p = _write(tmp_path, "app.py", "def hello():\n    return 'world'\n")
    r = peek(p)
    assert r.ext == ".py"
    assert r.extractable is True
    assert "def hello" in r.peek_text


def test_peek_csv_counts_rows(tmp_path: Path):
    body = "id,name,amount\n" + "\n".join(f"{i},row{i},{i*10}" for i in range(1, 51)) + "\n"
    p = _write(tmp_path, "sales.csv", body)
    r = peek(p)
    assert r.row_count == 50
    assert "row1" in r.peek_text


def test_peek_pdf_text_native(tmp_path: Path):
    pdf = _make_pdf(tmp_path, "doc.pdf", ["Hello world. " * 50] * 6)
    r = peek(pdf)
    assert r.page_count == 6
    assert r.extractable is True
    assert r.text_density >= 100


def test_peek_pdf_scanned_looks_sparse(tmp_path: Path):
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=2)
    r = peek(pdf)
    assert r.page_count == 2
    assert r.text_density < 50


def test_peek_docx_counts_images(tmp_path: Path):
    d = _make_docx(tmp_path, "mix.docx", paragraphs=5, images=3)
    r = peek(d)
    assert r.image_ratio > 0
    assert r.extractable is True


def test_peek_image_dims(tmp_path: Path):
    img = _make_image(tmp_path, "big.png", (800, 600))
    r = peek(img)
    assert r.image_dims == (800, 600)


def test_peek_unknown_extension(tmp_path: Path):
    p = _write(tmp_path, "weird.xyz", b"\x00\x01\x02")
    r = peek(p)
    assert r.peek_error is not None


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

def test_sensitivity_score_bank_statement(tmp_path: Path):
    p = _write(tmp_path, "stmt.txt", _sensitive_text())
    r = peek(p)
    assert compute_sensitivity_score(r) >= 7     # currency + totals + acct + dates


def test_sensitivity_score_neutral_text(tmp_path: Path):
    p = _write(tmp_path, "notes.txt", "The quick brown fox. Some musings about dogs.")
    r = peek(p)
    assert compute_sensitivity_score(r) == 0


def test_visual_score_scanned_pdf_high(tmp_path: Path):
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=2)
    r = peek(pdf)
    assert compute_visual_score(r) >= 7


def test_visual_score_text_pdf_low(tmp_path: Path):
    pdf = _make_pdf(tmp_path, "doc.pdf", ["Hello world. " * 100] * 10)
    r = peek(pdf)
    assert compute_visual_score(r) < 3


def test_visual_score_image_normal(tmp_path: Path):
    img = _make_image(tmp_path, "photo.png", (800, 600))
    r = peek(img)
    assert compute_visual_score(r) >= 4


def test_estimate_t4_cost_gpu_vs_cpu():
    # A 10-page file.
    class FakePeek:
        page_count = 10
    mb_gpu, s_gpu = estimate_t4_cost(FakePeek(), gpu_available=True)  # type: ignore[arg-type]
    mb_cpu, s_cpu = estimate_t4_cost(FakePeek(), gpu_available=False)  # type: ignore[arg-type]
    assert mb_gpu == mb_cpu == 2.0        # storage is hardware-independent
    assert s_gpu < s_cpu                   # CPU is slower


# ---------------------------------------------------------------------------
# decide() — per file type
# ---------------------------------------------------------------------------

def test_decide_text_small_is_t1(tmp_path: Path):
    p = _write(tmp_path, "note.md", "# hi\n")
    d = decide(peek(p))
    assert d.routes == ["T1"]
    assert d.skip_reason is None


def test_decide_text_huge_is_t0(tmp_path: Path):
    # 200 KB of text > 100 KB T0 threshold
    p = _write(tmp_path, "big.txt", "line\n" * 50_000)
    d = decide(peek(p))
    assert d.routes == ["T0"]


def test_decide_code_small_is_t1(tmp_path: Path):
    p = _write(tmp_path, "app.py", "def x(): pass\n")
    d = decide(peek(p))
    assert d.routes == ["T1"]


def test_decide_code_huge_is_t0(tmp_path: Path):
    p = _write(tmp_path, "big.js", "// x\n" * 200_000)
    d = decide(peek(p))
    assert d.routes == ["T0"]


def test_decide_csv_small_is_t1(tmp_path: Path):
    body = "a,b\n" + "\n".join(f"{i},{i}" for i in range(500))
    p = _write(tmp_path, "small.csv", body)
    d = decide(peek(p))
    assert d.routes == ["T1"]


def test_decide_csv_under_20mb_is_t1_rowlevel(tmp_path: Path):
    """CSV routing changed in b8ecab8 (row-level CSV indexing): anything
    under 20 MB goes T1 row-by-row; T2/T0 apply only above the size cutoff.
    This asserted the pre-row-level tiers until the 2026-08-30 triage."""
    body = "a,b\n" + "\n".join(f"{i},{i}" for i in range(5_000))
    p = _write(tmp_path, "mid.csv", body)
    d = decide(peek(p))
    assert d.routes == ["T1"]


def test_decide_csv_over_20mb_routes_by_rowcount(tmp_path: Path):
    """Above CSV_SIZE_T1_MAX (20 MB), rowcount picks T2 vs T0 (b8ecab8)."""
    # ~26 MB, ~1.2M rows: over the size cutoff AND over CSV_ROWS_T2_MAX
    body = "a,b\n" + "\n".join(f"{i},{i % 97},padpadpadpad" for i in range(1_200_000))
    p = _write(tmp_path, "huge.csv", body)
    assert p.stat().st_size > 20 * 1024 * 1024, "fixture must exceed the T1 size cap"
    d = decide(peek(p))
    assert d.routes == ["T0"]


def test_decide_short_pdf_is_t3(tmp_path: Path):
    pdf = _make_pdf(tmp_path, "short.pdf", ["Short PDF content."] * 3)
    d = decide(peek(pdf), gpu_available=True)
    assert d.routes == ["T3"]


def test_decide_long_text_pdf_is_t2(tmp_path: Path):
    # 40 pages, text-heavy, non-critical content → T2
    pdf = _make_pdf(tmp_path, "book.pdf", ["Chapter text. " * 200] * 40)
    d = decide(peek(pdf), gpu_available=True)
    assert d.routes == ["T2"]


def test_decide_scanned_pdf_gpu_is_t4(tmp_path: Path):
    # 10 empty pages — visual_score high, no text
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=10)
    d = decide(peek(pdf), gpu_available=True)
    assert d.routes == ["T4"]


def test_decide_scanned_pdf_no_gpu_falls_back_to_t3(tmp_path: Path):
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=10)
    d = decide(peek(pdf), gpu_available=False)
    # 10 pages * 10s/page = 100s > 10s CPU cap → fall back
    assert d.routes == ["T3"]


def test_decide_image_gpu_is_t4(tmp_path: Path):
    img = _make_image(tmp_path, "photo.jpg", (1200, 900))
    d = decide(peek(img), gpu_available=True)
    assert d.routes == ["T4"]


def test_decide_image_thumbnail_is_skipped(tmp_path: Path):
    img = _make_image(tmp_path, "tiny.png", (64, 64))
    # File bytes under threshold as well
    d = decide(peek(img), gpu_available=True)
    assert d.skipped
    assert d.skip_reason == "thumbnail"


def test_decide_docx_text_heavy_is_t2(tmp_path: Path):
    d = _make_docx(tmp_path, "notes.docx", paragraphs=20, images=0)
    dec = decide(peek(d), gpu_available=True)
    assert dec.routes == ["T2"]


def test_decide_docx_figure_heavy_gpu_is_t4(tmp_path: Path):
    d = _make_docx(tmp_path, "slides.docx", paragraphs=3, images=5)
    dec = decide(peek(d), gpu_available=True)
    assert dec.routes == ["T4"]


def test_decide_unknown_extension_is_skipped(tmp_path: Path):
    p = _write(tmp_path, "unknown.zzz", "whatever")
    d = decide(peek(p), gpu_available=True)
    assert d.skipped
    assert "unsupported" in (d.skip_reason or "")


# ---------------------------------------------------------------------------
# Additive criticality: T3 added on sensitive content
# ---------------------------------------------------------------------------

def test_decide_critical_pdf_gets_t3_plus_t2(tmp_path: Path):
    # 10-page PDF, text-native, with sensitive content on every sampled page.
    bank = _sensitive_text()
    pdf = _make_pdf(tmp_path, "stmt.pdf", [bank] * 10)
    d = decide(peek(pdf), gpu_available=True)
    assert "T3" in d.routes
    assert "T2" in d.routes
    assert d.criticality == "critical"
    assert d.criticality_source == "auto"


def test_decide_critical_scanned_pdf_gets_t3_plus_t4(tmp_path: Path):
    # Scanned PDF, critical via user override.
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=10)
    d = decide(peek(pdf), gpu_available=True, magpieconfig={"accuracy": "critical"})
    assert "T3" in d.routes
    assert "T4" in d.routes
    assert d.criticality == "critical"
    assert d.criticality_source == "user"


# ---------------------------------------------------------------------------
# T4 gates — per-file and corpus
# ---------------------------------------------------------------------------

def test_decide_t4_per_file_cap_gates_huge_scan(tmp_path: Path):
    # 300 pages × 0.2 MB = 60 MB > 50 MB per-file cap → falls back to T3.
    pdf = _make_image_pdf(tmp_path, "big_scan.pdf", n_pages=300)
    d = decide(peek(pdf), gpu_available=True)
    assert d.routes == ["T3"]
    assert any("over_per_file_cap" in n or "fell back" in n for n in d.notes)


def test_decide_t4_corpus_budget_exhausted(tmp_path: Path):
    pdf = _make_image_pdf(tmp_path, "scan.pdf", n_pages=10)
    # Simulate budget at 99% used
    d = decide(
        peek(pdf),
        gpu_available=True,
        t4_budget_used_mb=DEFAULT_T4_BUDGET_MB - 0.5,
    )
    assert d.routes == ["T3"]
    assert any("budget_exhausted" in n or "fell back" in n for n in d.notes)


# ---------------------------------------------------------------------------
# .magpieconfig.yaml (+ legacy .nasconfig.yaml)
# ---------------------------------------------------------------------------

def test_nasconfig_critical_override(tmp_path: Path):
    folder = tmp_path / "Bank Statements"
    folder.mkdir()
    (folder / ".nasconfig.yaml").write_text("accuracy: critical\n", encoding="utf-8")
    # A neutral text-native PDF inside this folder should be upgraded to critical.
    pdf = _make_pdf(folder, "ledger.pdf", ["ledger " * 100] * 10)
    cfg = load_magpieconfig(pdf)
    assert cfg.get("accuracy") == "critical"
    d = decide(peek(pdf), magpieconfig=cfg, gpu_available=True)
    assert "T3" in d.routes
    assert d.criticality_source == "user"


def test_nasconfig_colpali_never(tmp_path: Path):
    folder = tmp_path / "Private"
    folder.mkdir()
    (folder / ".nasconfig.yaml").write_text("colpali: never\n", encoding="utf-8")
    pdf = _make_image_pdf(folder, "scan.pdf", n_pages=10)
    cfg = load_magpieconfig(pdf)
    d = decide(peek(pdf), magpieconfig=cfg, gpu_available=True)
    assert d.routes == ["T3"]     # T4 disabled by config


def test_magpieconfig_preferred_name(tmp_path: Path):
    folder = tmp_path / "docs"
    folder.mkdir()
    (folder / ".magpieconfig.yaml").write_text("accuracy: critical\n", encoding="utf-8")
    pdf = folder / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    cfg = load_magpieconfig(pdf)
    assert cfg.get("accuracy") == "critical"


def test_magpieconfig_preferred_beats_legacy_in_same_folder(tmp_path: Path):
    folder = tmp_path / "docs"
    folder.mkdir()
    (folder / ".magpieconfig.yaml").write_text("accuracy: critical\n", encoding="utf-8")
    (folder / ".nasconfig.yaml").write_text("accuracy: casual\n", encoding="utf-8")
    pdf = folder / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    assert load_magpieconfig(pdf).get("accuracy") == "critical"


def test_nasconfig_missing_returns_empty(tmp_path: Path):
    p = _write(tmp_path, "x.txt", "hi")
    assert load_magpieconfig(p) == {}
